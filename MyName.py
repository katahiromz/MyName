#!/usr/bin/env python
# -*- coding: utf-8 -*-
# ワガナ（ワード画像並べ） by 片山博文MZ

# sys, os, datetime
import sys, os, datetime
import winreg as reg

class MyNameError(Exception):
    def __init__(self, msg):
        self.msg = msg
    def __str(self):
        return self.msg

# tkinter
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import messagebox
root = tk.Tk()

# Win32 API
import ctypes
from ctypes import c_long, WINFUNCTYPE
from ctypes.wintypes import HWND, UINT, WPARAM, LPARAM, DWORD
WM_DROPFILES = 0x0233
GWL_WNDPROC = -4
DragAcceptFiles = ctypes.windll.shell32.DragAcceptFiles
DragQueryFileW = ctypes.windll.shell32.DragQueryFileW
DragFinish = ctypes.windll.shell32.DragFinish
DragFinish.argtypes = [ ctypes.c_void_p ]
CallWindowProcW = ctypes.windll.user32.CallWindowProcW
CallWindowProcW.argtypes = [ ctypes.c_void_p, HWND, UINT, WPARAM, LPARAM ]
SetWindowLongW = ctypes.windll.user32.SetWindowLongW
GetModuleFileNameW = ctypes.windll.kernel32.GetModuleFileNameW

# window procedure
org_proc = None
dropped = []

@WINFUNCTYPE(c_long, HWND, UINT, WPARAM, LPARAM)
def win_proc(hwnd, msg, wp, lp):
    if msg == WM_DROPFILES:
        nf = DragQueryFileW(wp, -1, None, 0)
        for i in range(nf):
            buf = ctypes.create_unicode_buffer(260)
            if DragQueryFileW(wp, i, buf, 260):
                path = buf.value
                if os.path.isfile(path) and not os.path.isdir(path):
                    dropped.append(path)
        DragFinish(wp)
    return CallWindowProcW(org_proc, hwnd, msg, wp, lp)

# 定数。
NOSPEC = "(指定なし)"
COMPANY_KEY = "SOFTWARE\\Katayama Hirofumi MZ"
SOFT_KEY = COMPANY_KEY + "\\MyName"

# 空文字列ならNOSPECを返す。文字列を整える。
def NOSPEC_if_empty(text):
    text = text.strip()
    text.translate(str.maketrans({chr(0x0021 + i): chr(0xFF01 + i) for i in range(94)}))
    if text == "":
        return NOSPEC
    return text

# 切り詰める。
def truncate(string, length, ellipsis='...'):
    ret = ""
    i = 0
    for ch in string:
        if i + len(ellipsis) >= length:
            ret += ellipsis
            break;
        ret += ch
        if ord(ch) > 0x7F:
            i += 2
        else:
            i += 1
    return ret

# タイトルを変換する。
def convert_title(title, filename, image_index, the_time, char_limit=0):
    # %F、%N, %nの変換。
    filename = os.path.basename(filename)
    filename.replace("%", "%%")
    filename, ext = os.path.splitext(filename)
    title = title.replace("%F", filename)
    title = title.replace("%N", str(image_index + 1))
    title = title.replace("%n", str(image_index))
    # 日時書式を変換。
    title = the_time.strftime(title)
    import re
    title = re.sub(r'[\\/:*?"<>|]+', '_', title)
    # 長い場合は省略。
    if char_limit > 0:
        title = truncate(title, char_limit - 4, "[...]")
    return title

def try_int(value, field, another_value = "\x7F"):
    value = str(value).strip()
    value.translate(str.maketrans({chr(0x0021 + i): chr(0xFF01 + i) for i in range(94)}))
    if value == NOSPEC or value == str(another_value):
        return
    try:
        n = int(value)
    except:
        messagebox.showerror("ERROR", "「" + field + "」の欄が間違っています。")
        raise

current_filename = ""

class MyNameApplication(ttk.Frame):
    def reset_settings(self):
        self.muki_list = [NOSPEC, "縦向き", "横向き"]
        self.muki_default = NOSPEC
        self.gyousuu_list = ["1", "2", "3", "4"]
        self.gyousuu_default = "2"
        self.retsusuu_list = ["1", "2", "3", "4"]
        self.retsusuu_default = "2"
        self.page_title_default = "写真一覧(%N)"
        self.page_title_align_default = "中央揃え"
        self.page_size_list = [NOSPEC, "A4", "A3", "A2", "B5", "B4", "B3"]
        self.page_size_default = NOSPEC
        self.image_title_list = [NOSPEC, "画像 #%N", "%F", "(%N) %F", "%Y年%m月%d日", "%Y年%m月%d日 %H時%M分", "%Y-%m-%d", "%Y-%m-%d %H:%M"]
        self.image_title_default = "(%N) %F"
        self.file_title_list = ["写真一覧", "写真一覧(%Y年%m月%d日)", "私のアルバム(%Y年%m月%d日)", "実験結果-%Y-%m-%d", "押収品-%Y年%m月%d日"]
        self.file_title_default = "写真一覧"
        self.datetime_type_list = ["撮影日時", "画像作成日時", "画像更新日時", "ワード生成日時"]
        self.datetime_type_default = "ワード生成日時"
        self.auto_rotation_default = '1'
        self.button_advanced_hidden = False
    def __init__(self, root):
        super().__init__(root)
        self.image_ext_list = [".jpg", ".jpeg", ".jpe", ".jfif", ".png", ".gif", ".tif", ".tiff",
                               ".bmp", ".dib"]
        # リストと規定値。
        self.reset_settings()
        # レジストリから設定を読み込む。
        self.first_run = not self.load_settings()
        # フィルターを作成。
        self.filter = "*" + ";*".join(self.image_ext_list)
        # ウィジェットをすべて作成。
        self.create_widgets()
        self.pack()
        # アイコン設定。
        try:
            dir = os.path.dirname(os.getcwd() + "/" + __file__)
            file = dir + "\\icon.ico"
            if not os.path.isfile(file):
                dir = os.path.dirname(os.getcwd() + "/" + __file__)
                file = dir + "\\data\\icon.ico"
            if not os.path.isfile(file):
                buf = ctypes.create_unicode_buffer(260)
                GetModuleFileNameW(None, buf, 260)
                dir = os.path.dirname(buf.value)
                file = dir + "\\data\\icon.ico"
            if os.path.isfile(file):
                root.iconbitmap(file)
        except:
            pass
        # ウィンドウハンドルを取得。
        self.hwnd = self.winfo_id()
        # ドラッグ＆ドロップの準備。
        self.dnd_setup()
        self.dnd_interval = 600
        # 起動コマンドライン引数を処理。
        args = sys.argv[1:]
        self.dnd_notify(args)
        # 初回起動の場合は説明を表示。
        if (self.first_run):
            messagebox.showinfo("ワガナ",
                "「ワガナ」（ワード画像並べ）は、複数の画像を並べて、ワード文書ファイルにするソフトです。\n\n" +
                "基本操作は、画像ファイルをリストに追加して「ワード生成」ボタンを" +
                "押すだけです。完成したワードファイルはデスクトップに作成されます。")
    # ドラッグ＆ドロップされた。
    def dnd_notify(self, filenames):
        for filename in filenames:
            self.insert(filename)
    # ドラッグ＆ドロップを検査。
    def drop_check(self):
        global dropped
        if dropped:
            filenames = dropped
            dropped = []
            self.dnd_notify(filenames)
        self.after(self.dnd_interval, self.drop_check)
    # ドラッグ＆ドロップの準備。
    def dnd_setup(self):
        DragAcceptFiles(self.hwnd, True)
        global org_proc
        org_proc = SetWindowLongW(self.hwnd, GWL_WNDPROC, win_proc)
        self.after_idle(self.drop_check)
    # リストボックスの選択が変わった。
    def listbox_on_sel_change(self, evt=None):
        self.label_22.config(text="")
        selection = self.listbox_01.curselection()
        if len(selection) <= 0:
            self.button_04.config(state="disabled")
            self.label_18.image = None
            self.label_18["image"] = None
            return
        self.button_04.config(state="normal")
        filename = self.listbox_01.get(selection[0])
        picture_filename = self.process_image(filename, 50, 50)
        from PIL import Image, ImageTk
        img = Image.open(picture_filename)
        img = img.resize((50, 50))
        img = ImageTk.PhotoImage(img);
        self.label_18.image = img
        self.label_18["image"] = img
        the_time = self.get_datetime(filename)
        self.image_title_default = self.image_title.get()
        text = convert_title(self.image_title_default, filename, selection[0], the_time)
        self.label_22.config(text=("画像タイトル「" + text + "」"))
        if filename != picture_filename:
            os.remove(picture_filename)
    # リストボックスの選択が変わった。
    def combobox_on_sel_change(self, evt=None):
        self.listbox_on_sel_change()
    # ウィジェットをすべて作成。
    def create_widgets(self):
        self.group2 = tk.Frame(self)

        row = 0

        self.label_06 = ttk.Label(self.group2, text="ページサイズ:", width="", state="normal", )
        self.label_06.grid(column=0, row=row, pady=2, sticky=tk.E)
        self.page_size = tk.StringVar()
        self.combobox_06 = ttk.Combobox(self.group2, height="10", state="readonly", width="25", values=self.page_size_list, textvariable=self.page_size)
        self.combobox_06.grid(column=1, row=row, pady=2)
        self.combobox_06.set(self.page_size_default)
        row += 1

        self.label_01 = ttk.Label(self.group2, text="用紙の向き:", width="", state="normal", )
        self.label_01.grid(column=0, row=row, pady=2, sticky=tk.E)
        self.muki = tk.StringVar()
        self.combobox_01 = ttk.Combobox(self.group2, height="10", state="readonly", width="25", values=self.muki_list, textvariable=self.muki)
        self.combobox_01.grid(column=1, row=row, pady=2)
        self.combobox_01.set(self.muki_default)
        row += 1

        self.label_02 = ttk.Label(self.group2, text="ページのタテ割り:", width="", state="normal", )
        self.label_02.grid(column=0, row=row, pady=2, sticky=tk.E)
        self.gyousuu = tk.StringVar()
        self.combobox_02 = ttk.Combobox(self.group2, height="10", state="normal", width="25", values=self.gyousuu_list, textvariable=self.gyousuu)
        self.combobox_02.grid(column=1, row=row, pady=2)
        self.combobox_02.set(self.gyousuu_default)
        row += 1

        self.label_03 = ttk.Label(self.group2, text="ページのヨコ割り:", width="", state="normal", )
        self.label_03.grid(column=0, row=row, pady=2, sticky=tk.E)
        self.retsusuu = tk.StringVar()
        self.combobox_03 = ttk.Combobox(self.group2, height="10", state="normal", width="25", values=self.retsusuu_list, textvariable=self.retsusuu)
        self.combobox_03.grid(column=1, row=row, pady=2)
        self.combobox_03.set(self.retsusuu_default)
        row += 1

        row2 = 0
        self.label_17 = ttk.Label(self.group2, text="日時の種類:", width="", state="normal", )
        self.label_17.grid(column=2, row=row2, pady=2, sticky=tk.E)
        self.datetime_type = tk.StringVar()
        self.combobox_14 = ttk.Combobox(self.group2, height="10", state="readonly", width="25", values=self.datetime_type_list, textvariable=self.datetime_type)
        self.combobox_14.grid(column=3, row=row2, pady=2)
        self.combobox_14.set(self.datetime_type_default)
        self.combobox_14.bind('<<ComboboxSelected>>', self.combobox_on_sel_change)
        row2 += 1

        self.label_19 = ttk.Label(self.group2, text="画像タイトル:", width="", state="normal", )
        self.label_19.grid(column=2, row=row2, pady=2, sticky=tk.E)
        self.image_title = tk.StringVar()
        self.combobox_15 = ttk.Combobox(self.group2, height="10", state="normal", width="25", values=self.image_title_list, textvariable=self.image_title)
        self.combobox_15.grid(column=3, row=row2, pady=2)
        self.combobox_15.set(self.image_title_default)
        self.combobox_15.bind('<<ComboboxSelected>>', self.combobox_on_sel_change)
        row2 += 1

        self.label_20 = ttk.Label(self.group2, text="出力ファイル名:", width="", state="normal", )
        self.label_20.grid(column=2, row=row2, pady=2, sticky=tk.E)
        self.file_title = tk.StringVar()
        self.combobox_16 = ttk.Combobox(self.group2, height="10", state="normal", width="25", values=self.file_title_list, textvariable=self.file_title)
        self.combobox_16.grid(column=3, row=row2, pady=2)
        self.combobox_16.set(self.file_title_default)
        row2 += 1

        self.auto_rotation = tk.StringVar()
        self.checkbox_01 = ttk.Checkbutton(self.group2, text='余白が小さくなるように画像を回転', variable=self.auto_rotation)
        self.checkbox_01.grid(column=2, row=row2, columnspan=2)
        self.auto_rotation.set(self.auto_rotation_default)
        row2 += 1

        row += 1

        self.group4 = tk.Frame(self.group2)
        self.label_18 = ttk.Label(self.group4, width="50", state="normal", image="")
        self.label_18.pack(side=tk.LEFT, fill=tk.Y, padx=16)
        self.label_22 = ttk.Label(self.group4, width="", text="", state="normal")
        self.label_22.pack(side=tk.LEFT, fill=tk.Y, padx=16)
        self.label_21 = ttk.Label(self.group4, width="", text="全部で0個です。", state="normal")
        self.label_21.pack(side=tk.LEFT, fill=tk.Y, padx=16)

        self.group4.grid(column=0, row=row, columnspan=4)
        self.group2.rowconfigure(row, minsize=50)
        row += 1

        self.group1 = tk.Frame(self.group2)
        self.group1.grid(column=0, row=row, columnspan=4)
        row += 1

        self.label_16 = ttk.Label(self.group1, text="並べる画像ファイルのリスト: (ここに順番にファイルをドロップするか、画像データをCtrl+Vで貼り付けて下さい)", width="", state="normal", )
        self.label_16.pack(side=tk.TOP)

        self.listbox_01 = tk.Listbox(self.group1, width=92, height=10, selectmode=tk.EXTENDED, activestyle='none', exportselection=False)
        self.listbox_01.pack(side=tk.LEFT, fill=tk.Y)
        self.listbox_01.config(borderwidth=2)
        self.vscrollbar = tk.Scrollbar(self.group1, orient=tk.VERTICAL, command=self.listbox_01.yview)
        self.vscrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox_01.config(yscrollcommand = self.vscrollbar.set)
        self.listbox_01.bind('<<ListboxSelect>>', self.listbox_on_sel_change)

        self.group3 = tk.Frame(self.group2)
        self.group3.grid(column=0, row=row, columnspan=4, rowspan=1)
        row += 1

        self.button_01 = ttk.Button(self.group3, command = self.commandSaveList, text="リストの保存...", width="15", state="normal", )
        self.button_01.pack(side=tk.LEFT, padx=2)

        self.button_07 = ttk.Button(self.group3, command = self.commandResetSettings, text="設定のリセット", width="", state="normal", )
        self.button_07.pack(side=tk.LEFT, padx=2)

        self.button_02 = ttk.Button(self.group3, command = self.commandMoveUp, text="↑", width="8", state="normal", )
        self.button_02.pack(side=tk.LEFT, padx=2)

        self.button_03 = ttk.Button(self.group3, command = self.commandMoveDown, text="↓", width="8", state="normal", )
        self.button_03.pack(side=tk.LEFT, padx=2)

        self.button_04 = ttk.Button(self.group3, command = self.commandDeleteItems, text="選択を削除", width="15", state="normal", )
        self.button_04.pack(side=tk.LEFT, padx=2)

        self.button_05 = ttk.Button(self.group3, command = self.commandOK, text="ワード生成", width="", state="normal", )
        self.button_05.pack(side=tk.LEFT, padx=2)

        self.button_06 = ttk.Button(self.group3, command = self.commandExit, text="終了", width="", state="normal", )
        self.button_06.pack(side=tk.LEFT, padx=2)

        self.group2.columnconfigure(0, weight=1)
        self.group2.columnconfigure(1, weight=1)
        self.group2.columnconfigure(2, weight=1)
        self.group2.columnconfigure(3, weight=1)
        self.group2.pack()

        self.update_count()
    # 個数を更新。
    def update_count(self):
        if self.listbox_01.size() <= 0:
            self.button_04.config(state="disabled")
            return
        self.button_04.config(state="normal")
        self.label_21.config(text=("全部で" + str(self.listbox_01.size()) + "個です。"))
        self.listbox_on_sel_change()
    # 挿入。
    def insert(self, filename):
        if filename.endswith('.txt') or filename.endswith('.TXT'):
            try:
                lines = []
                with open(filename, "r") as fp:
                    lines = fp.readlines()
                self.listbox_01.delete(0, tk.END)
                for line in lines:
                    line = line.rstrip()
                    self.insert(line)
            except:
                messagebox.showerror('ワガナ', "ファイル「" + filename + "」の読み込みに失敗しました。")
            return
        self.listbox_01.selection_clear(0, tk.END)
        ext = os.path.splitext(filename)[1].lower()
        if ext in self.image_ext_list:
            self.listbox_01.insert(tk.END, filename)
            self.listbox_01.selection_set(self.listbox_01.size() - 1)
        self.update_count()
    # 「リストの保存」ボタンを押した。
    def commandSaveList(self):
        from tkinter import filedialog
        filename = filedialog.asksaveasfilename(initialdir=".", title="テキストファイルの保存先", \
            filetypes = (("テキストファイル", '*.txt'), ("すべてのファイル", "*.*")), \
            defaultextension=".txt")
        items = self.get_file_list()
        try:
            with open(filename, "w") as fp:
                for item in items:
                    fp.write(item + "\n")
            messagebox.showinfo('ワガナ', "ファイル「" + filename + "」の保存に成功しました。\n\nファイルドロップで復元できます。")
        except:
            messagebox.showerror('ワガナ', "ファイル「" + filename + "」の保存に失敗しました。")
    # 「選択を削除」ボタンを押した。
    def commandDeleteItems(self):
        items = self.listbox_01.curselection()
        for item in reversed(list(items)):
            self.listbox_01.delete(item)
        self.update_count();
    # 「↑」ボタンを押した。
    def commandMoveUp(self):
        if self.listbox_01.size() == 0:
            return
        items = list(self.listbox_01.curselection())
        items.sort()
        if items[0] == 0:
            return;
        for item in items:
            text0 = self.listbox_01.get(item - 1)
            text1 = self.listbox_01.get(item)
            self.listbox_01.delete(item - 1)
            self.listbox_01.delete(item - 1)
            self.listbox_01.insert(item - 1, text0)
            self.listbox_01.insert(item - 1, text1)
        for item in items:
            self.listbox_01.selection_clear(item)
        for item in items:
            self.listbox_01.selection_set(item - 1)
        self.listbox_on_sel_change()
    # 「↓」ボタンを押した。
    def commandMoveDown(self):
        if self.listbox_01.size() == 0:
            return
        items = list(self.listbox_01.curselection())
        items.sort()
        items.reverse()
        if items[0] == self.listbox_01.size() - 1:
            return;
        for item in items:
            text0 = self.listbox_01.get(item)
            text1 = self.listbox_01.get(item + 1)
            self.listbox_01.delete(item)
            self.listbox_01.delete(item)
            self.listbox_01.insert(item, text0)
            self.listbox_01.insert(item, text1)
        for item in items:
            self.listbox_01.selection_clear(item)
        for item in items:
            self.listbox_01.selection_set(item + 1)
        self.listbox_on_sel_change()
    # リストボックスからファイルのリストを取得する。
    def get_file_list(self):
        file_list = []
        size = self.listbox_01.size()
        for i in range(size):
            text = self.listbox_01.get(i)
            file_list.append(text)
        return file_list
    # 画像ファイルを処理する。処理後のファイル名を返す。
    def process_image(self, filename, contents_width, contents_height):
        from PIL import Image, ExifTags
        image = Image.open(filename)
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation] == 'Orientation':
                break
        exif = None
        try:
            exif = dict(image._getexif().items())
        except:
            pass
        # 必要なら回転する。
        if exif:
            if exif[orientation] == 3:
                image = image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image = image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image = image.rotate(90, expand=True)
        # 自動回転だったらさらに回転する。
        if self.auto_rotation.get() == '1':
            if image.width > image.height:
                if contents_width < contents_height:
                    image = image.rotate(-90, expand=True)
            elif image.width < image.height:
                if contents_width > contents_height:
                    image = image.rotate(-90, expand=True)
        # 一時ファイルに保存する。
        import tempfile
        ext = os.path.splitext(filename)[1].lower()
        handle, picture_filename = tempfile.mkstemp()
        os.close(handle)
        picture_filename += ext
        image.save(picture_filename)
        # ファイルサイズが1MB以上なら縮小する。
        while os.path.getsize(picture_filename) > 1024 * 1024:
            image = Image.open(picture_filename)
            new_image = image.resize((image.width // 2, image.height // 2))
            new_image.save(picture_filename)
        return picture_filename
    # 日時を取得する。
    def get_datetime(self, filename):
        the_time = None
        self.datetime_type_default = self.datetime_type.get()
        # 日時文字列の処理。
        if self.datetime_type_default == "撮影日時":
            try:
                from PIL import Image
                s = Image.open(filename)._getexif()[36867]
                s = s.replace(":", "-", 2)
                the_time = datetime.datetime.fromisoformat(s)
                #print("EXIF: " + filename + " | " + str(the_time))
            except:
                timestamp = os.path.getctime(filename)
                the_time = datetime.datetime.fromtimestamp(timestamp)
                #print("non-EXIF: " + filename + " | " + str(the_time))
        elif self.datetime_type_default == "画像作成日時":
            timestamp = os.path.getctime(filename)
            the_time = datetime.datetime.fromtimestamp(timestamp)
        elif self.datetime_type_default == "画像更新日時":
            timestamp = os.path.getmtime(filename)
            the_time = datetime.datetime.fromtimestamp(timestamp)
        elif self.datetime_type_default == "ワード生成日時":
            the_time = datetime.datetime.now()
        else:
            the_time = datetime.datetime.now()
        return the_time
    # docxファイルを生成する。
    def generate_docx(self):
        import docx
        from docx.table import Table
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Mm, Inches, Pt
        from docx.oxml.ns import qn

        # リストボックスからファイルのリストを取得する。
        file_list = self.get_file_list()

        # 最初の段落を取得する。
        dir = os.path.dirname(__file__)
        file = dir + "\\template.docx"
        if not os.path.isfile(file):
            dir = os.path.dirname(__file__)
            file = dir + "\\..\\template.docx"
        if not os.path.isfile(file):
            dir = os.path.dirname(__file__)
            file = dir + "\\..\\..\\template.docx"
        try:
            document = docx.Document(file)
        except:
            messagebox.showerror("ワガナ", "生成に必要なファイル「template.docx」が見つかりません。")
            return False

        para = document.paragraphs[0]

        #for style in document.styles:
        #    print(style)

        # フォント名。最初の段落を使う。
        the_font_name = para.style.font.name

        # フォントサイズ(mm)。
        if para.style.font.size == None:
            the_font_size = int(document.styles['Normal'].font.size.mm)
        else:
            the_font_size = int(para.style.font.size.mm)
        #print("the_font_size: " + str(the_font_size))

        # 文書をクリア。
        document._body.clear_content()

        # すべてのセクションにて。
        section = document.sections[-1]

        # ページサイズの指定を取得する(mm)。
        PAGE_SIZE_INFO = [
            ["A2", 420, 594],
            ["A3", 297, 420],
            ["A4", 210, 297],
            ["A5", 148, 210],
            ["B2", 515, 728],
            ["B3", 364, 515],
            ["B4", 257, 364],
            ["B5", 182, 257]]
        for info in PAGE_SIZE_INFO:
            if info[0] == self.page_size_default:
                section.page_width = Mm(info[1])
                section.page_height = Mm(info[2])
                break

        # 横向き指定なら横向きにする。
        page_width, page_height = section.page_width, section.page_height
        if self.muki_default == "横向き":
            section.orientation = WD_ORIENT.LANDSCAPE
            if page_width < page_height:
                page_width, page_height = page_height, page_width
        # 縦向き指定なら縦向きにする。
        if self.muki_default == "縦向き":
            section.orientation = WD_ORIENT.PORTRAIT
            new_width, new_height = section.page_height, section.page_width
            if page_height < page_width:
                page_width, page_height = page_height, page_width
        section.page_width = page_width
        section.page_height = page_height

        # 実際のページサイズ(mm)。
        the_page_height = section.page_height.mm
        the_page_width = section.page_width.mm

        # 余白を計算する(mm)。
        x_margin = section.left_margin.mm + section.right_margin.mm
        y_margin = section.top_margin.mm + section.bottom_margin.mm

        # 行数と列数。
        lines = int(self.gyousuu_default)
        columns = int(self.retsusuu_default)

        # 表の印刷可能領域(mm)。
        table_width = the_page_width - x_margin
        table_height = the_page_height - y_margin
        table_height -= the_font_size * 1.5
        if self.page_title_default != NOSPEC:
            table_height -= the_font_size

        # セルの最大サイズ(mm)。
        cell_max_width = int(table_width / columns)
        cell_max_height = int(table_height / lines)
        #print(cell_max_width)

        # セルの中身の最大サイズ(mm)。
        contents_width = table_width
        contents_height = table_height
        if self.image_title_default != NOSPEC:
            contents_height -= the_font_size * 2 * lines
        contents_width = int(contents_width / columns)
        contents_height = int(contents_height / lines)

        # 文字数制限。
        char_limit = int(contents_width * 1.7 / the_font_size)
        #print("char_limit: " + str(char_limit))

        # ページセル数。
        page_cells = lines * columns

        if len(file_list) <= 0:
            raise MyNameError("画像ファイルリストが空です。")

        # 実際に生成する。
        image_index = 0
        col_index = 0
        row_index = 0
        page_number = 0
        table = None

        for i in range(0, len(file_list)):
            filename = file_list[image_index]
            # 日時を取得する。
            the_time = self.get_datetime(filename)
            if i % page_cells == 0:
                # ページの最初のセル。
                if i != 0:
                    # 文書の最初のセルでなければ、直前に作成した行の高さを調整する。
                    for row in table.rows:
                        row.height = Mm(int(cell_max_height))
                    # ページ区切りを挿入。
                    document.add_page_break()
                # 必要ならば、ページ見出しを追加する。
                if self.page_title_default != NOSPEC:
                    # 段落を挿入する。
                    para = document.add_paragraph()
                    title = convert_title(self.page_title_default, "", page_number, the_time)
                    run = para.add_run(title)
                    # フォント設定。
                    run.font.name = the_font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), the_font_name)
                    run.font.size = Mm(the_font_size)
                    the_align = self.page_title_align_default
                    if the_align == "左揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif the_align == "中央揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif the_align == "右揃え":
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # 表を挿入する。行はゼロ個。
                table = document.add_table(rows=0, cols=columns)
                table.allow_autofit = False
                # 表を中央揃えにする。
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 行の直前なら
            if i % columns == 0:
                # 行を追加する。
                row_cells = table.add_row().cells
            # セルを取得し、縦に中央揃えする。
            cell = row_cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # セルの最初の段落を左右中央揃えにする。
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 段落に欄を追加。
            run = para.add_run()
            # 画像を処理する。
            requested_width = "max"
            requested_height = "max"
            global current_filename
            current_filename = filename
            picture_filename = self.process_image(filename, contents_width, contents_height)
            real_image_width = real_image_height = 0
            from PIL import Image
            with Image.open(picture_filename) as image:
                real_image_width = image.width
                real_image_height = image.height
            # 画像のサイズに応じて要求サイズを変更。
            if real_image_width * real_image_height != 0:
                aspect0 = float(contents_height) / float(contents_width)
                aspect1 = float(real_image_height) / float(real_image_width)
                if aspect0 < aspect1:
                    requested_height = contents_height
                    requested_width = requested_height / aspect1
                else:
                    requested_width = contents_width
                    requested_height = requested_width * aspect1
                requested_width *= 0.98
                requested_height *= 0.98
            # 画像を貼り付ける。
            if requested_width == NOSPEC and requested_height == NOSPEC:
                run.add_picture(picture_filename)
            elif requested_width != NOSPEC and requested_height == NOSPEC:
                run.add_picture(picture_filename, width = Mm(int(requested_width)))
            elif requested_width == NOSPEC and requested_height != NOSPEC:
                run.add_picture(picture_filename, height = Mm(int(requested_height)))
            elif requested_width != NOSPEC and requested_height != NOSPEC:
                run.add_picture(picture_filename, width = Mm(int(requested_width)), height = Mm(int(requested_height)))
            else:
                raise
            # 必要なら一時ファイルを削除。
            if filename != picture_filename:
                os.remove(picture_filename)
            # 画像のタイトルを取得する。
            image_title = self.image_title_default
            if image_title != NOSPEC:
                image_title = convert_title(image_title, filename, image_index, the_time, char_limit)
                # タイトル用の段落を追加する。
                para = cell.add_paragraph()
                # タイトルを入れる。
                run = para.add_run(image_title + "\n")
                # タイトルのフォント設定。
                run.font.name = the_font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), the_font_name)
                run.font.size = Mm(the_font_size)
                # 中央揃えにする。
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # インデックスを更新する。
            image_index += 1
            col_index += 1
            col_index %= columns
            if col_index == 0:
                row_index += 1
                if (image_index % page_cells == 0):
                    page_number += 1
        # もし、セルの高さが未指定でなければ
        for row in table.rows:
            row.height = Mm(int(cell_max_height))
        # 変更内容を保存する。
        title = convert_title(self.file_title_default, "", 0, the_time)
        title_with_ext = title + ".docx"
        dir = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        filename = dir + "/" + title_with_ext
        # すでに存在するか？
        if os.path.isfile(filename):
            answer = messagebox.askquestion(title='ワガナ', message="デスクトップに同名のファイル「" + title_with_ext + "」はすでに存在します。\n\n別名で保存しますか？", type=messagebox.YESNOCANCEL)
            if answer == messagebox.CANCEL:
                return False
            if answer == messagebox.YES:
                i = 0
                while os.path.isfile(filename):
                    i += 1
                    title_with_ext = title + "(" + str(i) + ")" + ".docx"
                    filename = dir + "/" + title_with_ext
        try:
            document.save(filename)
        except:
            raise MyNameError("ファイル「" + filename + "」はロックされていたため、保存に失敗しました。")
        messagebox.showinfo("ワガナ", "デスクトップにファイル「" + title_with_ext + "」の保存に成功しました。")
        return True
    # 設定を反映する。
    def apply(self):
        # 設定を取り出す。
        self.muki_default = self.muki.get()
        self.gyousuu_default = self.gyousuu.get()
        self.retsusuu_default = self.retsusuu.get()
        self.page_size_default = self.page_size.get()
        self.datetime_type_default = self.datetime_type.get()
        self.image_title_default = self.image_title.get()
        self.file_title_default = self.file_title.get()
        self.auto_rotation_default = self.auto_rotation.get()

        try_int(self.gyousuu_default, "行数")
        try_int(self.retsusuu_default, "列数")

        # 空文字列ならNOSPECにする。
        self.muki.set(NOSPEC_if_empty(self.muki_default))
        self.gyousuu.set(NOSPEC_if_empty(self.gyousuu_default))
        self.retsusuu.set(NOSPEC_if_empty(self.retsusuu_default))
        self.page_size.set(NOSPEC_if_empty(self.page_size_default))
        self.datetime_type.set(NOSPEC_if_empty(self.datetime_type_default))
        self.image_title.set(NOSPEC_if_empty(self.image_title_default))
        self.file_title.set(NOSPEC_if_empty(self.file_title_default))
        self.auto_rotation.set(self.auto_rotation_default)
    # 「生成する」ボタンを押した。
    def commandOK(self):
        # 設定を反映する。
        self.apply()

        # 設定内容を元に生成する。
        try:
            self.generate_docx()
        except Exception as err:
            global current_filename
            if current_filename != "":
                messagebox.showerror("ワガナ", "ファイル「" + current_filename + "」の処理に失敗しました。")
            else:
                messagebox.showerror("ワガナ", "何らかの処理に失敗しました。")
            return
        # 設定を保存する。
        self.save_settings()
    # 「終了」ボタンを押した。
    def commandExit(self):
        # 設定を反映する。
        self.apply()
        # 設定を保存する。
        self.save_settings()
        # GUIを破棄する。
        root.destroy()
    # 「設定の初期化」ボタンを押した。
    def commandResetSettings(self):
        self.reset_settings()
        self.save_settings()
        self.muki.set(self.muki_default)
        self.gyousuu.set(self.gyousuu_default)
        self.retsusuu.set(self.retsusuu_default)
        self.page_size.set(self.page_size_default)
        self.datetime_type.set(self.datetime_type_default)
        self.image_title.set(self.image_title_default)
        self.file_title.set(self.file_title_default)
        self.auto_rotation.set(self.auto_rotation_default)
        self.listbox_on_sel_change()
        messagebox.showinfo("ワガナ", "設定を初期化しました。")
    # １つ設定を読み込む。
    def read_settings(self, key, name, the_list, value):
        try:
            count = int(reg.QueryValueEx(key, name + "_count")[0])
            if count >= 0:
                the_list.clear()
                for i in range(count):
                    value = reg.QueryValueEx(key, name + "_" + str(i))[0]
                    the_list.append(value)
            return reg.QueryValueEx(key, name + "_value")[0]
        except:
            return value
    # １つ設定を書き込む。
    def write_settings(self, key, name, the_list, value):
        if value not in the_list:
            the_list.append(value)
        reg.SetValueEx(key, name + "_value", 0, reg.REG_SZ, value)
        reg.SetValueEx(key, name + "_count", 0, reg.REG_SZ, str(len(the_list)))
        for i, item in enumerate(the_list):
            reg.SetValueEx(key, name + "_" + str(i), 0, reg.REG_SZ, item)
        return True
    # レジストリから設定を読み込む。
    def load_settings(self):
        first_run = True
        try:
            with reg.OpenKeyEx(reg.HKEY_CURRENT_USER, SOFT_KEY, 0, reg.KEY_READ|reg.KEY_WOW64_64KEY) as soft_key:
                first_run = False
                self.muki_default = self.read_settings(soft_key, "muki", self.muki_list, self.muki_default)
                self.gyousuu_default = self.read_settings(soft_key, "gyousuu", self.gyousuu_list, self.gyousuu_default)
                self.retsusuu_default = self.read_settings(soft_key, "retsusuu", self.retsusuu_list, self.retsusuu_default)
                self.page_size_default = self.read_settings(soft_key, "page_size", self.page_size_list, self.page_size_default)
                self.datetime_type_default = self.read_settings(soft_key, "datetime_type", self.datetime_type_list, self.datetime_type_default)
                self.image_title_default = self.read_settings(soft_key, "image_title", self.image_title_list, self.image_title_default)
                self.file_title_default = self.read_settings(soft_key, "file_title", self.file_title_list, self.file_title_default)
                self.auto_rotation_default = reg.QueryValueEx(soft_key, "auto_rotation_value")[0] == 'yes'
        except:
            pass
        return not first_run
    # レジストリに設定を書き込む。
    def save_settings(self):
        try:
            with reg.CreateKeyEx(reg.HKEY_CURRENT_USER, SOFT_KEY, 0, reg.KEY_WRITE|reg.KEY_WOW64_64KEY) as soft_key:
                self.write_settings(soft_key, "muki", self.muki_list, self.muki_default)
                self.write_settings(soft_key, "gyousuu", self.gyousuu_list, self.gyousuu_default)
                self.write_settings(soft_key, "retsusuu", self.retsusuu_list, self.retsusuu_default)
                self.write_settings(soft_key, "page_size", self.page_size_list, self.page_size_default)
                self.write_settings(soft_key, "datetime_type", self.datetime_type_list, self.datetime_type_default)
                self.write_settings(soft_key, "image_title", self.image_title_list, self.image_title_default)
                self.write_settings(soft_key, "file_title", self.file_title_list, self.file_title_default)
                reg.SetValueEx(soft_key, 'auto_rotation_value', 0, reg.REG_SZ, 'yes' if self.auto_rotation_default == '1' else 'no')
                return True
        except:
            return False

# 主処理。
root.title('ワガナ（ワード画像並べ） Version 1.2 by 片山博文MZ')

# サイズを変更させない。
root.resizable(False, False)

# メインアイコンを設定。
try:
    root.iconbitmap('./data/icon.ico')
except:
    root.iconbitmap('./icon.ico')

# メインウィンドウを作成。
frame = MyNameApplication(root)

# 一時ファイルの個数を表す変数。
iImage = 0

# ウィンドウを閉じたときの処理。
def on_closing():
    # 一時ファイルを削除する。
    global iImage
    while iImage >= 0:
        try:
            fname = os.getenv('TMP') + "\\Image-" + str(iImage) + ".png"
            os.remove(fname)
            #print(fname)
        except:
            pass
        iImage -= 1
    # 終了。
    frame.commandExit()
root.protocol("WM_DELETE_WINDOW", on_closing)

# Ctrl+Vを押したときの処理。
def on_ctrl_v(e):
    from PIL import Image, ImageGrab
    try:
        # クリップボードにある画像を一時ファイルに保存する。
        clip_img = ImageGrab.grabclipboard()
        if clip_img:
            global iImage
            iImage += 1
            fname = os.getenv('TMP') + "\\Image-" + str(iImage) + ".png"
            clip_img.save(fname, "PNG")
            frame.insert(fname)
            #print(fname)
    except:
        pass
root.bind('<Control-v>', on_ctrl_v)

root.mainloop()
